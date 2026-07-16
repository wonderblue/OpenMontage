import os
import sys
import re
from fountain import Fountain
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Add scripts directory to path to import main screenplay strings
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from build_docx_via_fountain import VIDEO_1_FOUNTAIN, VIDEO_2_FOUNTAIN, VIDEO_3_FOUNTAIN, set_font

# Define the visual storyboard descriptions for each scene
STORYBOARD_DATA = {
    "Video_1_Foundations_of_Marketing": {
        "1.1": "VISUAL: Medium Close-Up of the Professor avatar standing in a modern university library study, with warm lighting and soft background bokeh.\n\nGRAPHIC: Slide 1 is displayed on a virtual screen: \"UNIT-1: Importance and Scope of Marketing\".\n\nTEXT OVERLAY: \"Welcome to Marketing Management\"",
        "2.1": "VISUAL: Scene cuts to a vibrant, sunny aerial view of the Guntur Chilli Market, crowded with traders, farmers, and piles of red chillies.\n\nGRAPHIC: On-screen list slides in:\n* 1. Market Research\n* 2. Strategy\n* 3. Branding\n\nTEXT OVERLAY: \"Guntur Sannam Chilli - GI Tagged Brand\"",
        "2.2": "VISUAL: Split screen. Left side: A busy Rythu Bazar showing direct transaction between farmers and consumers in Vijayawada. Right side: A retail textile shop in Brodipet with a clear discount pricing tag (999/-).\n\nGRAPHIC: A beautiful travel poster mockup of AP Tourism beach locations (Visakhapatnam, Amaravati) slides into view.",
        "2.3": "VISUAL: Scene changes to a restaurant in Arundelpet showing dinner combo layouts.\n\nGRAPHIC: A digital news article overlay slides in: \"Guntur Medical College Achieves National Excellence Ranking\". A Gantt chart play animation demonstrates the steps of the Amaravati development project.",
        "3.1": "VISUAL: Splits showing traditional handloom weavers operating looms in Mangalagiri (Production Concept), transitioning to a detailed close-up of an artisan hand-painting wooden toys in Kondapalli (Product Concept).\n\nGRAPHIC: Word overlay: \"Marketing Myopia: Losing Sight of the Customer\".",
        "3.2": "VISUAL: An animated scale slides in, showing an aggressive salesman pushing a box (Selling) versus a student-centric course curriculum planner (Marketing).\n\nGRAPHIC: Transition to lush green fields of Andhra Pradesh with the text overlay: \"Zero Budget Natural Farming (APCNF)\" (Societal Concept).",
        "4.1": "VISUAL: An interactive circular wheel chart of the \"8 Tasks of Marketing Management\" rotates on screen, lighting up each task (insights, connection, brands, offerings, growth) as the narrator speaks.",
        "5.1": "VISUAL: A nested concentric ring animation expands outwards:\n* Inner Ring: Internal Environment (Faculty Culture, Assets)\n* Middle Ring: Micro Environment (Suppliers, Competitors)\n* Outer Ring: Macro Environment (GST, Tech Mandis, Weather)",
        "6.1": "VISUAL: A mathematical equation scale balances: Customer Value = Total Benefits - Total Costs.\n\nGRAPHIC: Side-by-side comparison icons of local Guntur Main Bazaar personal trust vs. corporate Smart Bazaar convenience.",
        "7.1": "VISUAL: Split screen comparison: An individual buying a phone (B2C) vs. a formal business boardroom where managers evaluate quality certificates from Visakhapatnam Steel Plant (Vizag Steel) and Godrej Agrovet (B2B).",
        "8.1": "VISUAL: Return to the library. Medium Close-up of Professor avatar.\n\nGRAPHIC: Slide 1 checklist slides down, marking all foundations topics as completed."
    },
    "Video_2_Services_Marketing": {
        "1.1": "VISUAL: Medium Close-up of Professor avatar standing in the library.\n\nGRAPHIC: Slide 9 title card slides onto screen: \"UNIT-1: Services Marketing\".",
        "2.1": "VISUAL: A dynamic smartphone mockup slides onto screen. The screen animates through the Uber passenger app booking route, driver rating cards (4.9 vs. 4.2), and safety sharing map routes in Guntur/Vijayawada.",
        "3.1": "VISUAL: Infographic icons appear for service characteristics:\n* Intangibility: Icon of a cloud storage promise.\n* Perishability: APSRTC bus departing, clock ticking down.\n* Inseparability: A classroom screen splitting to show both teacher and student actively engaging (Co-creation).",
        "3.2": "VISUAL: Detailed animations:\n* Variability: A hot, steaming plate of Guntur Biryani with spice level icons shifting depending on chef moods.\n* Changing Demand: Araku Valley tourist temperature and booking volume chart showing winter peaks vs. summer valleys.",
        "4.1": "VISUAL: A 3D honeycomb diagram of the 7Ps of Services appears. Product, Price, Place, and Promotion cells expand.\n\nGRAPHIC: Stethoscope icons (KIMS healthcare product), subscription postpaid packages (Airtel pricing), Swiggy delivery route (Place), and student success posters (Promotion).",
        "4.2": "VISUAL: The remaining 7Ps cells expand: People, Process, Physical Evidence.\n\nGRAPHIC: Front-desk receptionist at Novotel Visakhapatnam (People), Mee Seva digital workflow steps (Process), and modern air-conditioned coaching lobby infrastructure in Guntur (Physical Evidence).",
        "5.1": "VISUAL: Returns to Professor in the library, recapping services and highlighting the transition to global markets."
    },
    "Video_3_Global_Marketing": {
        "1.1": "VISUAL: Medium Close-up of Professor avatar welcome in the library.\n\nGRAPHIC: Slide 14 title card appears: \"UNIT-1: Global Marketing\".",
        "2.1": "VISUAL: A global map with trading routes connecting continents. A balancing scale compares \"Standardization\" vs. \"Adaptation\".\n\nGRAPHIC: A comic illustration of the baby food jar packaging misunderstanding in Africa.",
        "3.1": "VISUAL: Splits showing:\n* Nike: US basketball sponsorship shifting to European soccer, then Indian Cricket blue jersey. Zoom into specialized spikes on cricket shoes.\n* Domino's: American pizza menu morphing into Paneer Tikka pizza (India) and squid toppings (Japan). Delivery rider on a motorcycle navigates traffic.",
        "4.1": "VISUAL: Grid layout showing the first four challenges:\n* Cultural: McDonald's menu showing McAloo Tikki and Maharaja Mac.\n* Legal: Checklist showing GDPR logo and cosmetic safety testing symbols.\n* Economic: Sachet packaging designs.\n* Supply Chain: Logistics cold-chain refrigeration containers.",
        "4.2": "VISUAL: Grid layout showing the remaining four challenges:\n* Political: Trade tariff percentages overlay.\n* Language: Blunder examples text (KFC \"Eat your fingers off\" / GM Nova \"It doesn't go\").\n* Pricing/HR: Gray market textbook shipment routing, and US top-down vs. Japanese bottom-up Ringi consensus charts.",
        "5.1": "VISUAL: Final library scene. Complete unit timeline checklist checkmarks. Professor bids farewell."
    }
}

def clean_script_text(text):
    # Remove TECHNICAL NOTES and visual transition blocks to keep it clean
    text = re.sub(r"^>.*$", "", text, flags=re.MULTILINE)
    text = re.sub(r"TECHNICAL NOTES:\s*(?:-\s*.*\n?)*", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text

def convert_to_storyboard_docx(filename, raw_script, storyboard_map, output_path):
    # Process Fountain parser
    fountain_parser = Fountain()
    tokens = fountain_parser.tokenize(raw_script.splitlines())
    
    doc = Document()
    
    # Page setup - Margins (wider to accommodate two-column table)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Title/Metadata Header
    title_keys = ["subject", "title", "page length", "duration"]
    title_dict = {}
    
    for token in tokens:
        if token.get("type") == "titles":
            title_dict = token.get("title-fields", {})
            break
            
    if title_dict:
        for k in title_keys:
            val = title_dict.get(k)
            if val:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                k_cap = " ".join([w.capitalize() for w in k.split()])
                run = p.add_run(f"{k_cap}: {val}")
                set_font(run, bold=True)
        doc.add_paragraph()
        
    current_scene_num = None
    table = None
    left_cell = None
    right_cell = None
    
    for token in tokens:
        ttype = token.get("type")
        lines = token.get("lines", [])
        
        if not lines:
            continue
            
        text = "".join(lines).strip()
        
        if ttype == "titles":
            continue
            
        elif ttype == "slugline":
            # Extract scene number from heading if exists, e.g. ".SCENE 2.1"
            scene_match = re.search(r"SCENE\s+(\d+\.\d+)", text)
            if scene_match:
                current_scene_num = scene_match.group(1)
            else:
                current_scene_num = None
                
            # Create a new Table for this Scene Heading and its body content
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            
            # Widths: Left (Audio/Dialogue) = 3.5 inches, Right (Visual Storyboard) = 3.0 inches
            table.columns[0].width = Inches(3.5)
            table.columns[1].width = Inches(3.0)
            
            row = table.rows[0]
            left_cell = row.cells[0]
            right_cell = row.cells[1]
            
            left_cell.width = Inches(3.5)
            right_cell.width = Inches(3.0)
            
            # Add Scene Heading to Left Cell
            p = left_cell.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(text)
            set_font(run, bold=True)
            
            # Add Visual Storyboard content to Right Cell
            v_para = right_cell.add_paragraph()
            v_para.paragraph_format.space_before = Pt(6)
            v_para.paragraph_format.space_after = Pt(6)
            
            # Fetch storyboard description
            visual_desc = ""
            if current_scene_num and current_scene_num in storyboard_map:
                visual_desc = storyboard_map[current_scene_num]
            else:
                visual_desc = "VISUAL: Professor speaking to camera."
                
            # Write visual description (right-aligned lines)
            v_run = v_para.add_run(visual_desc)
            set_font(v_run, size=10, italic=True)
            
        elif ttype == "dialogue":
            if left_cell is not None:
                # Character name
                char_p = left_cell.add_paragraph()
                char_p.paragraph_format.space_before = Pt(6)
                char_p.paragraph_format.space_after = Pt(0)
                run_char = char_p.add_run(token.get("character", ""))
                set_font(run_char, bold=True)
                
                # Dialogue lines
                dialogue_lines = token.get("text", [])
                for line in dialogue_lines:
                    if line.startswith("(") and line.endswith(")"):
                        p_para = left_cell.add_paragraph()
                        p_para.paragraph_format.space_after = Pt(0)
                        run_p = p_para.add_run(line)
                        set_font(run_p, italic=True)
                    else:
                        d_para = left_cell.add_paragraph()
                        d_para.paragraph_format.space_after = Pt(6)
                        run_d = d_para.add_run(line)
                        set_font(run_d)
                        
        elif ttype == "transition":
            # Skip transition block inside the table to keep it clean
            continue
            
        elif ttype == "action":
            # Audio descriptions or actions
            if left_cell is not None:
                p = left_cell.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(text)
                if text.startswith("AUDIO:"):
                    set_font(run, size=11, bold=True, italic=True)
                else:
                    set_font(run, size=11, italic=True)
                    
        # Add a spacing paragraph after each scene table
        doc.add_paragraph()
        
    doc.save(output_path)
    print(f"Exported storyboard script to Word Doc: {output_path}")

def generate_storyboards():
    storyboard_dir = "/Users/rkanduk/Documents/GitHub/OpenMontage/storyboard"
    if not os.path.exists(storyboard_dir):
        os.makedirs(storyboard_dir)
        print(f"Created folder: {storyboard_dir}")
        
    scripts = {
        "Video_1_Foundations_of_Marketing": (VIDEO_1_FOUNTAIN, STORYBOARD_DATA["Video_1_Foundations_of_Marketing"]),
        "Video_2_Services_Marketing": (VIDEO_2_FOUNTAIN, STORYBOARD_DATA["Video_2_Services_Marketing"]),
        "Video_3_Global_Marketing": (VIDEO_3_FOUNTAIN, STORYBOARD_DATA["Video_3_Global_Marketing"])
    }
    
    for filename, (raw_script, storyboard_map) in scripts.items():
        # Clean script
        clean_script = clean_script_text(raw_script)
        
        # Save Fountain file
        fountain_path = os.path.join(storyboard_dir, f"{filename}.fountain")
        with open(fountain_path, "w", encoding="utf-8") as f:
            # For Fountain file, we save the clean script
            f.write(clean_script)
        print(f"Saved Storyboard Fountain script: {fountain_path}")
        
        # Export Docx with Split Storyboard Table
        docx_path = os.path.join(storyboard_dir, f"{filename}.docx")
        convert_to_storyboard_docx(filename, clean_script, storyboard_map, docx_path)

if __name__ == "__main__":
    generate_storyboards()
