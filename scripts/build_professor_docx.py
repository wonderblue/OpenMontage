import os
import sys
import re
from fountain import Fountain
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Add scripts directory to path to import main screenplay strings
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from build_docx_via_fountain import VIDEO_1_FOUNTAIN, VIDEO_2_FOUNTAIN, VIDEO_3_FOUNTAIN, set_font

def strip_technical_details(script_text):
    text = re.sub(r"Background:.*", "Background:", script_text)
    text = re.sub(r"Voice:.*", "Voice:", text)
    text = re.sub(r"Sound Design:.*", "Sound Design:", text)
    text = re.sub(r"Sound Effects:.*", "Sound Effects:", text)
    text = re.sub(r"^>.*$", "", text, flags=re.MULTILINE)
    text = re.sub(r"TECHNICAL NOTES:\s*(?:-\s*.*\n?)*", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text

def add_page_number(doc):
    """Add page numbers to the header (top-right, screenplay standard)."""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Add page number field
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run._r.append(instrText)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)
        
        run.font.name = "Courier New"
        run.font.size = Pt(10)

def add_horizontal_rule(doc):
    """Add a thin horizontal line as a visual scene separator."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')  # thin line
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')  # light grey
    pBdr.append(bottom)
    pPr.append(pBdr)

# Map scene numbers to readable location captions for the storyboard image
SCENE_CAPTIONS = {
    "Video_1": {
        "1.1": "Modern University Library [Slide 1]",
        "2.1": "Guntur Chilli Market [Slide 1]",
        "2.2": "Brodipet Textile Market / Rythu Bazar [Slide 1]",
        "2.3": "Arundelpet Restaurant / Amaravati [Slide 1]",
        "3.1": "Mangalagiri Handloom / Kondapalli Workshop [Slides 2-3]",
        "3.2": "Guntur Educational Campus / Natural Farm [Slides 2-3]",
        "4.1": "Business Hub / APSRTC Booking [Slide 4]",
        "5.1": "Spice Trading Office / Supply Chain [Slide 5]",
        "6.1": "Guntur Main Bazar / Modern Retail [Slide 6]",
        "7.1": "Vizag Steel Plant / Agricultural Cooperative [Slide 7]",
        "8.1": "Modern University Library [Slide 1]"
    },
    "Video_2": {
        "1.1": "Modern University Library [Slide 1]",
        "2.1": "Streets of Vijayawada / Guntur [Slide 9]",
        "3.1": "KIMS Hospital / AP Tourism [Slide 11]",
        "3.2": "Mee Seva Center / Araku Valley [Slide 11]",
        "4.1": "Airtel Office / Coaching Center [Slides 12-13]",
        "4.2": "Novotel Hotel / Mee Seva Center [Slides 12-13]",
        "5.1": "Modern University Library [Slide 1]"
    },
    "Video_3": {
        "1.1": "Modern University Library [Slide 1]",
        "2.1": "Global Connectivity Map [Slide 14]",
        "3.1": "Retail Stores — Nike & Domino's [Slides 15-16]",
        "4.1": "Corporate HQ / Legal Office [Slide 19]",
        "4.2": "Global Operations / HR Center [Slide 20]",
        "5.1": "Modern University Library [Slide 1]"
    }
}

def convert_to_professor_final(filename, clean_script, image_folder, output_path):
    fountain_parser = Fountain()
    tokens = fountain_parser.tokenize(clean_script.splitlines())
    
    doc = Document()
    
    # Strict Screenplay Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1)
        
    # FIX 7: Add page numbers (top-right)
    add_page_number(doc)
        
    # Title / Metadata Header
    title_keys = ["subject", "title", "page length", "duration"]
    title_dict = {}
    
    for token in tokens:
        if token.get("type") == "titles":
            title_dict = token.get("title-fields", {})
            break
            
    # FIX 6: Center-align the metadata block
    if title_dict:
        for k in title_keys:
            val = title_dict.get(k)
            if val:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(6)
                k_cap = " ".join([w.capitalize() for w in k.split()])
                run = p.add_run(f"{k_cap}: {val}")
                set_font(run, name="Courier New", size=12, bold=True)
        doc.add_paragraph()
        
    current_scene_num = None
    video_key = ""
    if "Video_1" in filename:
        video_key = "Video_1"
    elif "Video_2" in filename:
        video_key = "Video_2"
    elif "Video_3" in filename:
        video_key = "Video_3"
    
    is_first_scene = True
        
    for token in tokens:
        ttype = token.get("type")
        lines = token.get("lines", [])
        
        if not lines:
            continue
            
        text = "".join(lines).strip()
        
        if ttype == "titles":
            continue
            
        elif ttype == "slugline":
            scene_match = re.search(r"SCENE\s+(\d+\.\d+)", text)
            if scene_match:
                current_scene_num = scene_match.group(1)
            else:
                current_scene_num = None
            
            # FIX 4: Add horizontal rule before each scene (except the first)
            if not is_first_scene:
                add_horizontal_rule(doc)
            is_first_scene = False
                
            # Scene heading
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(text.upper())
            set_font(run, name="Courier New", size=12, bold=True)
            
            # Storyboard image (compact 3.2 inch width)
            if current_scene_num and video_key:
                img_name = f"{video_key}_Scene_{current_scene_num}_Slide.png"
                img_path = os.path.join(image_folder, img_name)
                if os.path.exists(img_path):
                    v_para = doc.add_paragraph()
                    v_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    v_para.paragraph_format.space_before = Pt(4)
                    v_para.paragraph_format.space_after = Pt(2)
                    v_run = v_para.add_run()
                    v_run.add_picture(img_path, width=Inches(3.2))
                    
                    # FIX 5: Add image caption
                    caption_text = ""
                    captions = SCENE_CAPTIONS.get(video_key, {})
                    caption_text = captions.get(current_scene_num, f"Scene {current_scene_num}")
                    
                    cap_para = doc.add_paragraph()
                    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_para.paragraph_format.space_before = Pt(0)
                    cap_para.paragraph_format.space_after = Pt(12)
                    cap_run = cap_para.add_run(f"Scene {current_scene_num} — {caption_text}")
                    cap_run.font.name = "Courier New"
                    cap_run.font.size = Pt(9)
                    cap_run.italic = True
                    cap_run.font.color.rgb = RGBColor(128, 128, 128)
                else:
                    v_para = doc.add_paragraph()
                    v_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    v_para.paragraph_format.space_before = Pt(6)
                    v_para.paragraph_format.space_after = Pt(12)
                    v_run = v_para.add_run(f"[STORYBOARD: {img_name}]")
                    set_font(v_run, name="Courier New", size=10, italic=True)
                    
        elif ttype == "dialogue":
            # Character Name (centered per screenplay standard)
            char_p = doc.add_paragraph()
            char_p.paragraph_format.left_indent = Inches(2.2)
            char_p.paragraph_format.space_before = Pt(12)
            char_p.paragraph_format.space_after = Pt(0)
            run_char = char_p.add_run(token.get("character", "").upper())
            set_font(run_char, name="Courier New", size=12, bold=True)
            
            dialogue_lines = token.get("text", [])
            for line in dialogue_lines:
                if line.startswith("(") and line.endswith(")"):
                    p_para = doc.add_paragraph()
                    p_para.paragraph_format.left_indent = Inches(1.6)
                    p_para.paragraph_format.space_after = Pt(0)
                    run_p = p_para.add_run(line)
                    set_font(run_p, name="Courier New", size=12, italic=True)
                else:
                    # FIX 3: Reduced right indent from 1.5 to 0.75 inches
                    d_para = doc.add_paragraph()
                    d_para.paragraph_format.left_indent = Inches(1.0)
                    d_para.paragraph_format.right_indent = Inches(0.75)
                    d_para.paragraph_format.space_after = Pt(6)
                    run_d = d_para.add_run(line)
                    set_font(run_d, name="Courier New", size=12)
                    
        elif ttype == "action":
            # FIX 1: Detect INT./EXT. location lines and format as bold sluglines
            if text.startswith("INT.") or text.startswith("EXT."):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(12)
                run = p.add_run(text.upper())
                set_font(run, name="Courier New", size=12, bold=True)
            # FIX 2: Collapse AUDIO blocks into a compact single-line label
            elif text.startswith("AUDIO:"):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run("[Audio Cues]")
                run.font.name = "Courier New"
                run.font.size = Pt(10)
                run.italic = True
                run.font.color.rgb = RGBColor(153, 153, 153)
            elif text in ("Background:", "Voice:", "Sound Design:", "Sound Effects:"):
                # Skip the individual empty audio label lines
                continue
            else:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0)
                p.paragraph_format.space_after = Pt(12)
                run = p.add_run(text)
                set_font(run, name="Courier New", size=12, italic=True)
                
    doc.save(output_path)
    print(f"Exported final screenplay to Word Doc: {output_path}")

def generate_professor_version():
    professor_dir = "/Users/rkanduk/Documents/GitHub/OpenMontage/professor"
    storyboard_assets_dir = "/Users/rkanduk/Documents/GitHub/OpenMontage/professor_storyboard"
    
    if not os.path.exists(professor_dir):
        os.makedirs(professor_dir)
        print(f"Created folder: {professor_dir}")
        
    scripts = {
        "Video_1_Foundations_of_Marketing": (VIDEO_1_FOUNTAIN, os.path.join(storyboard_assets_dir, "Video_1_Foundations")),
        "Video_2_Services_Marketing": (VIDEO_2_FOUNTAIN, os.path.join(storyboard_assets_dir, "Video_2_Services")),
        "Video_3_Global_Marketing": (VIDEO_3_FOUNTAIN, os.path.join(storyboard_assets_dir, "Video_3_Global"))
    }
    
    for filename, (raw_content, img_folder) in scripts.items():
        clean_content = strip_technical_details(raw_content)
        
        # Save clean Fountain file
        fountain_path = os.path.join(professor_dir, f"{filename}.fountain")
        with open(fountain_path, "w", encoding="utf-8") as f:
            f.write(clean_content)
        print(f"Saved Professor Fountain script: {fountain_path}")
        
        # Convert and Save Docx
        docx_path = os.path.join(professor_dir, f"{filename}.docx")
        convert_to_professor_final(filename, clean_content, img_folder, docx_path)

if __name__ == "__main__":
    generate_professor_version()
