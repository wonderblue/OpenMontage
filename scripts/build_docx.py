import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_font(run, name="Courier New", size=12, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def format_docx(input_txt, output_docx):
    doc = Document()
    
    # Configure margins for standard screenplay look
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1)
        
    with open(input_txt, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    in_production_notes = False
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
            
        # Metadata / Title
        if line.startswith("Title:") or line.startswith("Credit:") or line.startswith("Duration:") or line.startswith("Format:"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(lines[i].strip())
            set_font(run, bold=True)
            i += 1
            continue
            
        # Section dividers
        if "====" in line:
            doc.add_page_break()
            i += 1
            continue
            
        # Production notes header
        if line.startswith("PRODUCTION NOTES:"):
            in_production_notes = True
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line)
            set_font(run, bold=True, size=14)
            i += 1
            continue
            
        if in_production_notes:
            # Check if production notes have ended (reached a scene heading)
            if line.startswith("SCENE"):
                in_production_notes = False
                # fall through to scene processing
            else:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(lines[i].strip())
                set_font(run, size=11, italic=True)
                i += 1
                continue

        # Scene heading
        if line.startswith("SCENE ") or line.startswith("INT.") or line.startswith("EXT."):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(lines[i].strip())
            set_font(run, bold=True)
            i += 1
            continue
            
        # Transitions / Technical Notes
        if line.startswith("TRANSITION:") or line.startswith("TECHNICAL NOTES:"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(3.5)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(lines[i].strip())
            set_font(run, italic=True)
            i += 1
            continue
            
        # Character & Dialogue detection
        # Speakers are all-caps words like PROFESSOR, FEMALE NARRATOR, MALE NARRATOR
        is_character = line in ["PROFESSOR", "FEMALE NARRATOR", "MALE NARRATOR"]
        if is_character:
            # Character name
            char_p = doc.add_paragraph()
            char_p.paragraph_format.left_indent = Inches(2.2)
            char_p.paragraph_format.space_before = Pt(12)
            char_p.paragraph_format.space_after = Pt(0)
            run = char_p.add_run(line)
            set_font(run, bold=True)
            
            # Read subsequent line(s) as dialogue
            i += 1
            dialogue_text = []
            while i < len(lines):
                next_line = lines[i].strip()
                if not next_line:
                    i += 1
                    continue
                # If next line is a scene heading, transition, character, or section divider, break
                if (next_line.startswith("SCENE ") or 
                    next_line.startswith("INT.") or 
                    next_line.startswith("EXT.") or 
                    next_line.startswith("TRANSITION:") or 
                    next_line.startswith("TECHNICAL NOTES:") or 
                    next_line in ["PROFESSOR", "FEMALE NARRATOR", "MALE NARRATOR"] or 
                    "====" in next_line):
                    break
                dialogue_text.append(lines[i].strip())
                i += 1
                
            if dialogue_text:
                full_dialogue = " ".join(dialogue_text)
                
                # Check for parenthetical at start
                m = re.match(r"^(\([^\)]+\))\s*(.*)", full_dialogue)
                if m:
                    parenthetical = m.group(1)
                    dialogue_content = m.group(2)
                    
                    # Add parenthetical paragraph
                    p_para = doc.add_paragraph()
                    p_para.paragraph_format.left_indent = Inches(1.6)
                    p_para.paragraph_format.space_after = Pt(0)
                    run_p = p_para.add_run(parenthetical)
                    set_font(run_p, italic=True)
                    
                    # Add dialogue paragraph
                    d_para = doc.add_paragraph()
                    d_para.paragraph_format.left_indent = Inches(1.0)
                    d_para.paragraph_format.right_margin = Inches(1.5)
                    d_para.paragraph_format.space_after = Pt(12)
                    run_d = d_para.add_run(dialogue_content)
                    set_font(run_d)
                else:
                    d_para = doc.add_paragraph()
                    d_para.paragraph_format.left_indent = Inches(1.0)
                    d_para.paragraph_format.right_margin = Inches(1.5)
                    d_para.paragraph_format.space_after = Pt(12)
                    run_d = d_para.add_run(full_dialogue)
                    set_font(run_d)
            continue

        # Standard Action / Visual description / Audio info
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        if line.startswith("AUDIO:") or line.startswith("VISUAL:") or line.startswith("Visual shows"):
            run = p.add_run(lines[i].strip())
            set_font(run, size=11, italic=True)
        else:
            run = p.add_run(lines[i].strip())
            set_font(run)
        i += 1
        
    doc.save(output_docx)
    print(f"Created Word Document at {output_docx}")

if __name__ == "__main__":
    base_dir = "/Users/rkanduk/Documents/GitHub/OpenMontage"
    format_docx(
        os.path.join(base_dir, "Video_1_Foundations_of_Marketing_Screenplay.txt"),
        os.path.join(base_dir, "Video_1_Foundations_of_Marketing_Screenplay.docx")
    )
    format_docx(
        os.path.join(base_dir, "Video_2_Advanced_Marketing_Concepts_Screenplay.txt"),
        os.path.join(base_dir, "Video_2_Advanced_Marketing_Concepts_Screenplay.docx")
    )
