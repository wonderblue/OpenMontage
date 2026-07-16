import os
import re
from fountain import Fountain
from docx import Document
from docx.shared import Inches, Pt

# Define Video 1 in Fountain format: Foundations of Marketing (Slides 1-8)
VIDEO_1_FOUNTAIN = """Subject: Marketing Management
Title: Foundations of Marketing
Page Length: Estimated ~21 pages
Duration: Around ~20 Minutes

.SCENE 1.1 - 00:00-01:30 [Slide 1]

INT. MODERN UNIVERSITY LIBRARY - DAY

AUDIO:
Background: Gentle Indian classical instrumental (sitar/flute fusion), low volume.
Voice: Professor's cloned voice (warm, authoritative, welcoming).
Sound Design: None (clean dialogue).

PROFESSOR
Namaste and welcome! I'm Professor [Name], and I'm delighted to introduce you to this comprehensive module on Marketing Management.

Marketing is the heartbeat of every successful business. It's not just about selling products—it's about understanding human needs, creating value, and building lasting relationships.

In this first video, we'll explore the fundamental importance and scope of marketing, understand the five evolutionary concepts that have shaped modern marketing, examine critical management tasks, analyze the complex marketing environment, and dive deep into customer value and industrial B2B marketing.

By the end of this journey, you'll have a solid foundation to understand how companies and services from Guntur to our broader regional markets have mastered the art and science of local marketing.

Let's begin this exciting exploration together.

TECHNICAL NOTES:
- Ensure perfect lip-sync with cloned voice
- Professional lighting on avatar (three-point lighting setup)
- Maintain eye contact with camera
- Background: AI-generated study with subtle depth of field

> Smooth fade to black (1 second), then fade into animated marketing ecosystem.

.SCENE 2.1 - 01:30-04:30 [Slide 1]

EXT. GUNTUR CHILLI MARKET - DAY

AUDIO:
Background: Light, upbeat instrumental (non-distracting).
Voice: Female narrator (warm, confident, storytelling tone).
Sound Effects: Subtle soft "whoosh" as each element appears.

FEMALE NARRATOR
Imagine marketing as the sun in a solar system. Around it orbit nine critical planets—each essential, each unique, each contributing to the success of the entire ecosystem.

These nine elements work together seamlessly to create marketing magic. Let's explore each one.

First: Market Research. Before a single rupee is spent, smart organizations listen. Market research is the compass that guides every decision. Take the Guntur Chillies market, for example. Before the Rabi season, traders and farmers analyze weather patterns, global demand from countries like Bangladesh and Sri Lanka, domestic consumption trends, and pricing patterns from previous years. This isn't guessing; it's science.

Second: Strategy. Research without strategy is like having a map but no destination. Strategy is the blueprint that turns insights into action. It answers the critical questions: Who are we targeting? What value do we offer? How do we differentiate ourselves?

Third: Branding. Your brand is your promise. It's what people say about you when you're not in the room. Think of Guntur Sannam chillies—the GI tag, the reputation for heat and quality, the trust built over generations. That's the power of strategic branding.

.SCENE 2.2 - 04:30-07:30 [Slide 1]

INT. BRODIPET TEXTILE MARKET / RYTHU BAZAR - DAY

AUDIO:
Background: Continuing light instrumental.
Voice: Female narrator (continuing).
Sound Effects: Soft transition sounds between visual elements.

FEMALE NARRATOR
Fourth: Distribution. How do you get your product from farm or factory to customer? This is where distribution strategy becomes critical. Consider the Rythu Bazars across Andhra Pradesh—direct farmer-to-consumer markets that eliminate middlemen, reduce costs, and ensure fresh produce reaches consumers in Vijayawada, Guntur, and Tenali. This is distribution excellence.

Fifth: Pricing. Pricing isn't just about numbers—it's psychology. It's perception. In the textile markets of Guntur's Brodipet, shopkeepers use psychological pricing—999 instead of 1000, seasonal discounts during Sankranti, bulk purchase incentives. They understand that pricing influences buyer behavior.

Sixth: Advertising. This is your megaphone to the world. But in today's cluttered media landscape, shouting louder doesn't work. You need to shout smarter. Consider how Andhra Pradesh Tourism promotes Amaravati, the Buddhist circuit, or beach tourism in Visakhapatnam. They don't just tell you destinations exist; they make you feel the experience, they create desire, they make you remember. That's effective advertising.

.SCENE 2.3 - 07:30-10:00 [Slide 1]

INT. ARUNDELPET RESTAURANT / AMARAVATI TIMELINE - DAY

AUDIO:
Background: Continuing instrumental, slightly more dynamic.
Voice: Female narrator (continuing).
Sound Effects: Soft notification sounds for social media elements.

FEMALE NARRATOR
Seventh: Promotion. While advertising is paid media, promotion encompasses the broader mix—sales promotions, discounts, loyalty programs, contests. Consider how local restaurants in Guntur's Arundelpet or Pattabhipuram offer lunch combos, festival discounts during Ugadi, or loyalty cards for regular customers. This is promotional strategy in action.

Eighth: Publicity. This is earned media—when others talk about you. It's more credible than advertising because it's not paid for. When Guntur Medical College gets featured in national rankings, or when a local startup from NTR District gets covered, that's publicity. It's organic, it's authentic, and it builds trust.

Ninth: Projects. Marketing isn't a one-time activity; it's a series of strategic projects. Launching Amaravati as a capital city, promoting Guntur as an educational hub, developing the Krishna River pushkarams—each requires careful planning, execution, and measurement.

These nine elements don't work in isolation. They're interconnected. A pricing decision affects your branding. Your distribution strategy impacts your promotion. Market research informs everything.

This is the scope of marketing—complex, dynamic, and absolutely critical to business success.

> Smooth fade to next segment with 1-second fade effect.

.SCENE 3.1 - 10:00-12:00 [Slides 2-3]

INT. MANGALAGIRI HANDLOOM / KONDAPALLI WORKSHOP - DAY

AUDIO:
Background: Slightly more dynamic instrumental (subtle tempo increase).
Voice: Female narrator (continuing, slightly more energetic).
Sound Effects: Clock ticking for timeline transitions.

FEMALE NARRATOR
Marketing hasn't always been customer-centric. It has evolved through five distinct concepts, each reflecting the business realities of its time.

Concept One: The Production Concept.

This is the oldest philosophy. It's based on a simple assumption: customers want products that are affordable and available everywhere.

In Andhra Pradesh, consider the handloom cooperatives in Mangalagiri. By focusing on efficient production and wide distribution through government emporiums and online platforms, they made traditional Mangalagiri sarees accessible and affordable across the state and beyond.

But here's the limitation: This concept works only when demand exceeds supply. In today's competitive market, just being cheap and available isn't enough.

Concept Two: The Product Concept.

This philosophy assumes customers will favor products that offer the most quality, performance, and innovative features.

Consider Kondapalli toys—the intricate craftsmanship, the vibrant colors, the cultural authenticity. Artisans don't just make toys; they create pieces of art that tell stories. The focus is on superior quality and unique features.

But there's a trap here—'Marketing Myopia.' You can become so obsessed with your product that you forget what business you're really in. If Kondapalli toy makers only focused on traditional designs without adapting to modern tastes, they would lose relevance. Innovation within tradition is key.

.SCENE 3.2 - 12:00-14:00 [Slides 2-3]

INT. GUNTUR EDUCATIONAL CAMPUS / NATURAL FARM - DAY

AUDIO:
Background: Continuing instrumental.
Voice: Female narrator (continuing, slightly more authoritative).
Sound Effects: Subtle transition sounds between concepts.

FEMALE NARRATOR
Concept Three: The Selling Concept.

This philosophy assumes that customers won't buy enough of your product unless you aggressively promote and sell it.

This is common with unsought goods—things people don't normally think of buying. The focus is on selling what the company makes, not making what the market wants.

The problem? It assumes customers who are coerced into buying will like the product. And even if they don't, they won't badmouth it. In the age of social media and online reviews, that's a dangerous assumption.

Concept Four: The Marketing Concept.

This is where modern marketing begins. The marketing concept flips the script. Instead of 'make and sell,' it's 'sense and respond.'

The customer is at the center of everything. All organizational activities—product development, pricing, distribution, promotion—are designed around creating customer value and satisfaction.

Consider how successful educational institutions in Guntur operate. They don't just offer courses; they research what skills employers need, what students want to learn, and design programs accordingly. They provide placement support, industry connections, and continuous feedback mechanisms. That's the marketing concept in action.

Concept Five: The Societal Marketing Concept.

This is the most advanced form. It asks: Is satisfying individual customer wants always good for society in the long run?

The societal marketing concept balances three considerations:
1. Company profits
2. Consumer wants
3. Society's interests

Consider the Zero Budget Natural Farming movement in Andhra Pradesh. It's not just about profitable farming or meeting consumer demand for food. It's about environmental sustainability, farmer welfare, soil health, and long-term societal benefit. This is marketing with a conscience.

Another example: How temples like Tirumala manage prasadam distribution, crowd control, and devotee services. They balance revenue generation with spiritual service and social responsibility.

These five concepts show the evolution of marketing thinking—from production efficiency to customer obsession to societal welfare. Understanding this evolution is crucial for modern marketers.

> Fade to next segment with subtle sound effect.

.SCENE 4.1 - 14:00-15:30 [Slide 4]

INT. BUSINESS HUB / APSRTC BOOKING DEPT - DAY

AUDIO:
Background: Steady, professional instrumental (mellow).
Voice: Female narrator (clear, confident).
Sound Effects: Soft notification for each task completion.

FEMALE NARRATOR
Marketing management isn't a single activity. It's a complex orchestration of eight critical tasks that successful organizations execute flawlessly.

Task One: Capture Marketing Insights.
This is about understanding the marketplace. Through market research, data analytics, and customer feedback, organizations gather intelligence. Consider how the Andhra Pradesh government conducted extensive research before implementing the Rythu Bharosa scheme or how Guntur traders analyze global spice market trends before setting auction prices.

Task Two: Connect with Customers.
This goes beyond transactions. It's about building relationships. Consider how local businesses in Guntur's main bazaar maintain customer relationships through personalized service, festival greetings, and understanding family needs across generations.

Task Three: Build Strong Brands.
A strong brand is a valuable asset. It commands premium prices, creates loyalty, and provides competitive advantage. Building a brand requires consistency, authenticity, and time. Think of how Guntur chillies, Tirupati laddu, or Machilipatnam Kalamkari have built brands over centuries.

Task Four: Shape Market Offerings.
This is where you design products and services that create value. It involves decisions about features, quality, design, branding, and packaging. Consider how Kondapalli toy artisans are now creating modern designs alongside traditional ones, or how AP tourism packages are being designed for different traveler segments.

Task Five: Deliver Value.
This is about making your offering available and accessible. Distribution channels, logistics, inventory management—all play a role. The success of e-Choupal in rural Andhra Pradesh or the efficiency of APSRTC's online booking system demonstrates delivery excellence.

Task Six: Communicate Value.
Even the best product fails if no one knows about it. This task involves advertising, public relations, sales promotion, personal selling, and digital marketing. It's about telling your story in a way that resonates. Consider how successful local brands in Vijayawada and Guntur use a mix of traditional media, social media, and word-of-mouth.

Task Seven: Create Long-Term Growth.
Sustainable marketing isn't about quarterly results; it's about long-term value creation. This requires innovation, market expansion, customer retention, and adaptation to changing environments. Look at how institutions like Siddhartha Academy or Vasavi Group have grown over decades.

Task Eight: Develop Market Strategies and Plans.
This is the strategic framework that ties everything together. It involves segmentation, targeting, positioning, and the development of comprehensive marketing plans. The Amaravati development masterplan is an example of strategic marketing at a macro level.

> Smooth fade to next segment (1 second).

.SCENE 5.1 - 15:30-17:00 [Slide 5]

INT. SPICE TRADING OFFICE / SUPPLY CHAIN - DAY

AUDIO:
Background: Thoughtful, slightly serious instrumental.
Voice: Male narrator (analytical, measured, slightly deeper tone).
Sound Effects: Subtle ambient sounds for each layer.

MALE NARRATOR
No organization operates in a vacuum. Every marketing decision is influenced by a complex web of environmental factors.

Think of it as three concentric circles.

The Inner Circle: The Internal Environment.
This is your organization—your culture, your resources, your capabilities, your people. Before you can serve customers, you must align your internal environment. Consider how successful educational institutions in Guntur build strong faculty cultures, invest in infrastructure, and create systems that enable excellence.

The Middle Circle: The Micro Environment.
These are the actors close to the organization that affect its ability to serve customers:

Suppliers: They provide the resources you need. For a Guntur chilli trader, suppliers are the farmers from surrounding villages. A disruption in this supply chain affects everything.

Distributors and Intermediaries: They're your bridge to customers. Commission agents in Guntur market, retail partners, online platforms—choose them wisely.

Customers: The center of everything. Understand their needs, their behaviors, their preferences. Are they domestic buyers, exporters, or end consumers?

Competitors: They're fighting for the same customers. Know their strengths, their weaknesses, their strategies. Other chilli markets like Byadgi or Teja compete with Guntur.

The Outer Circle: The Macro Environment.
These are larger societal forces that affect the entire microenvironment:

Technological: Digital mandis, online trading platforms, blockchain for traceability—technology is reshaping traditional markets.

Economic: Inflation, interest rates, economic growth, GST policies—all impact purchasing power and business decisions.

Environmental: Climate change affecting crop patterns, water scarcity, sustainability requirements—these are no longer optional considerations.

Social-Cultural: Demographics, values, lifestyles, festival demands, health consciousness—these shape what customers want.

Successful organizations don't just react to these environmental forces; they anticipate them, adapt to them, and sometimes even shape them.

Understanding this environment is not a one-time exercise. It's continuous. The market doesn't stand still, and neither can you.

> Smooth transition to next segment.

.SCENE 6.1 - 17:00-18:30 [Slide 6]

INT. GUNTUR MAIN BAZAR / MODERN RETAIL STORE - DAY

AUDIO:
Background: Light, engaging instrumental.
Voice: Female narrator (conversational, relatable, warm).
Sound Effects: Gentle scale balancing sounds.

FEMALE NARRATOR
Let's move to a fundamental concept: Customer Value.

Customer value is a customer's perception of what they get versus what they give. It's not just about price; it's about the total package of benefits.

Think of it as an equation:

Customer Value = Total Benefits - Total Costs

Benefits include:
- Service Quality: How well does the product or service perform?
- Social Value: What does using this brand say about me?
- Past Experience: How have previous interactions been?
- Emotional Satisfaction: How does this make me feel?

Costs include:
- Price: The monetary cost
- Time: How long does it take?
- Effort: How much work is required?
- Psychological Cost: Is there stress or risk involved?

Why do people choose to shop at Reliance Smart Bazaar or other modern supermarkets in Guntur instead of local kirana stores? The benefits—variety under one roof, air-conditioned comfort, parking facilities, promotional offers, branded products—create superior customer value for their target market.

Conversely, why do millions still prefer local markets like Guntur Main Bazaar or Brodipet? For them, the value equation is different. Personal relationships, credit facilities, bargaining power, neighborhood convenience, and trust built over years outweigh modern amenities.

Consider Tirupati Tirumala Devasthanam. Devotees don't just buy prasadam; they buy spiritual satisfaction, divine connection, and social recognition. The value is transcendental.

Understanding customer value is the foundation of all marketing strategy.

.SCENE 7.1 - 18:30-20:00 [Slide 7]

INT. VIZAG STEEL PLANT / AGRICULTURAL COOPERATIVE - DAY

AUDIO:
Background: Professional, steady instrumental (slightly corporate feel).
Voice: Male narrator (authoritative, B2B tone, clear).
Sound Effects: Business-like sounds, subtle document shuffling.

MALE NARRATOR
So far, we've focused on marketing to individual consumers. But there's another massive world: Industrial Marketing, also known as Business-to-Business or B2B marketing.

It's the creation and management of mutually beneficial relationships between organizational suppliers and organizational customers.

Your customer isn't an individual; it's a private firm, a public agency, or a nonprofit organization.

Let's compare:
B2C: You buy a smartphone for personal use. Decision: Quick. Emotional. Individual.
B2B: A construction company in Vijayawada buys cement for a building project. Decision: Slow. Rational. Multiple stakeholders. Procurement team. Technical evaluation. Finance approval. Vendor comparison. Negotiations. Contracts.

The Characteristics of Industrial Marketing:

1. Fewer but Larger Customers: Instead of millions of consumers, you might have hundreds of corporate clients. But each relationship is worth significantly more.

2. Complex Decision-Making: Multiple people are involved—users, influencers, decision-makers, gatekeepers. You're not selling to one person; you're selling to a committee.

3. Derived Demand: The demand for industrial products is derived from the demand for consumer products. Cement manufacturers sell to builders because people buy homes.

4. Professional Purchasing: B2B buyers are trained professionals. They evaluate specifications, compare vendors, negotiate terms. Data and value rule.

5. Long-Term Relationships: B2B is about partnerships that last years, even decades.

Benefits of Industrial Marketing:
- Improved Customer Experience: Deep understanding of client needs
- Precise Specification Delivery: Exact requirements met consistently
- Enhanced Lead Generation: Qualified prospects, higher conversion
- Long-Term Competitive Advantage: Relationships competitors can't easily break

Consider Visakhapatnam Steel Plant (Vizag Steel). They market to construction companies, infrastructure developers, and manufacturing units. Their marketing focuses on quality standards, delivery reliability, technical specifications, and long-term supply contracts.

Or consider Godrej Agrovet selling agricultural inputs to farmer cooperatives across Andhra Pradesh. Their marketing emphasizes product efficacy, technical support, training programs, and partnership.

B2B marketing is relationship marketing at its most strategic.

> Fade to professor outro (1 second).

.SCENE 8.1 - 20:00-21:30 [Slide 1]

INT. MODERN UNIVERSITY LIBRARY - DAY

AUDIO:
Background: Same Indian classical instrumental as intro.
Voice: Professor's cloned voice (warm, encouraging).
Sound Effects: Subtle text reveal sounds.

PROFESSOR
Excellent work! You've now completed the first video of our Marketing Management journey.

Let's recap what we've covered:
We've explored the scope of marketing, the five evolutionary concepts that have shaped modern marketing thinking, the critical tasks of marketing management, the internal, micro and macro environment, the fundamentals of customer value, and the unique dynamics of B2B industrial marketing.

We've seen how these concepts apply right here in our region—from Guntur chilli markets to Mangalagiri handlooms, from Kondapalli toys to Vizag Steel and local businesses.

This foundation is crucial. But we're just getting started.

In our next video, we will explore Services Marketing in detail, analyzing the unique characteristics that make services different, like Uber and local services, and studying the extended 7Ps marketing mix.

Join me in Video 2, where we'll continue this exciting exploration.

Namaste!

TECHNICAL NOTES:
- Maintain consistency with intro segment
- Ensure smooth audio transition
- Professional closing animation
"""

# Define Video 2 in Fountain format: Services Marketing (Slides 9-13) - Expanded Classroom Style
VIDEO_2_FOUNTAIN = """Subject: Marketing Management
Title: Services Marketing
Page Length: Estimated ~20 pages
Duration: Around ~20 Minutes

.SCENE 1.1 - 00:00-01:30 [Slide 1]

INT. MODERN UNIVERSITY LIBRARY - DAY

AUDIO:
Background: Indian classical instrumental (same as Video 1).
Voice: Professor's cloned voice (warm, welcoming back).
Sound Effects: None (clean dialogue).

PROFESSOR
Welcome back! I am so glad you've returned to continue our Marketing Management journey. 

In Video 1, we spent a lot of time talking about tangible items. We talked about Guntur chillies, textiles in Brodipet, and steel from Vizag. These are physical things. You can drop them on your foot, and it hurts. That makes it easy to understand. It has weight, dimensions, and you can see it and inspect its quality before spending a single rupee.

But today, we are going to look at something much more fascinating, and frankly, a bit trickier: Services Marketing. 

Services dominate our modern economy. In fact, if you look at the GDP of most developing and developed countries, services represent well over fifty percent of the total output. Think about the apps on your phone, the banking facilities you use, the schools we go to, and the healthcare we receive. 

In this video, we're going to dive headfirst into what makes service marketing completely different from product marketing. We will look at a famous case study, Uber, to understand how services build trust. Then, we will break down the six unique characteristics of services, and finally, look at why the traditional four Ps of marketing just aren't enough, which is why we must expand them to the seven Ps.

So, let's step into the world of the intangible. Sit back, take notes, and let's get started.

TECHNICAL NOTES:
- Maintain visual consistency with Video 1
- Reference previous video naturally
- Set clear expectations for this video

> Smooth fade to first content segment.

.SCENE 2.1 - 01:30-05:30 [Slide 9]

INT. STREETS OF VIJAYAWADA/GUNTUR - DAY

AUDIO:
Background: Energetic, modern instrumental.
Voice: Female narrator (engaging, story-driven).
Sound Effects: App notification sounds, car sounds.

FEMALE NARRATOR
Let's kick things off with a simple definition. What exactly is service marketing? 

At its core, service marketing is the process of promoting and selling an intangible transaction or experience to a target audience. You are selling a process or a performance rather than a physical object. 

Think about this: when you buy a smartphone, you walk out of the store with a box. You own that phone. You can show it off. But when you buy a hair styling service, or book a flight ticket, or get a medical check-up, what do you actually own at the end of the day? Nothing physical. You own the memory, the benefit, or the state of well-being. You are buying an experience.

Now let's talk about one of the most successful service brands in the world: Uber. 

If we look at Slide 9, Uber is an on-demand transportation technology that connects riders with driver-partners. But here is the impromptu classroom question: Is Uber really just selling a car ride? 

No. Absolutely not. If they were just selling car rides, they would be no different from a regular taxi. What Uber is actually marketing is convenience, flexibility, and most importantly, a temporary "trust architecture."

Think about the sheer audacity of the business model. You open an app on your phone, you request a ride, a stranger pulls up in a private car, you get into the back seat, and they drive you to your destination. Ten years ago, if your parents saw you doing this, they would have panicked. How did Uber make this normal? 

They built credibility and trust. 

Uber builds credibility using a double-sided rating system. After every ride, the driver rates the rider, and the rider rates the driver. This gamifies good behavior. It ensures that bad actors are weeded out. 

Let's go impromptu here. Have you ever checked your driver's rating and hesitated because it was a 4.2 instead of a 4.9? That rating tells you a story about service quality before you even open the car door. It's a psychological guarantee. If it's a 4.2, you wonder what went wrong. Did they drive too fast? Were they rude? That score is the brand's reputation condensed into a single digit.

Uber builds trust by implementing strict safety protocols—GPS tracking of every ride, driver background checks, in-app emergency buttons, and incident response teams. For drivers, they market flexibility—being their own boss. For riders, they market peace of mind. You can share your live location with your family. 

So, when you see an Uber driving down the streets of Vijayawada or Guntur, remember that they aren't just selling transport. They are selling trust. And in services marketing, trust is the ultimate currency.

.SCENE 3.1 - 05:30-10:30 [Slide 11]

INT. KIMS HOSPITAL / AP TOURISM DESK - DAY

AUDIO:
Background: Continuing light instrumental.
Voice: Female narrator (continuing).
Sound Effects: Subtle transition sounds.

FEMALE NARRATOR
Now, let's turn to Slide 11. To truly master service marketing, you must understand the six fundamental characteristics that make services unique. 

The first characteristic is Intangibility. Services cannot be seen, tasted, felt, or heard before they are purchased. 

Imagine you are going to KIMS Hospital in Guntur. You cannot "sample" the surgery before you buy it. You cannot test-drive the doctor's diagnosis. Because of this intangibility, customers face high levels of uncertainty. 

Here's a quick insight: we value what we can see. If an IT consultant fixes your network in two minutes by clicking one button, you sometimes feel cheated because you didn't see them work hard. But you are paying for twenty years of experience, not the two minutes. That is the paradox of intangibility. Customers want to see the labor.

So, how do marketers solve this? We must "tangibilize the intangible." We use physical cues—clean marble floors, professional uniforms, high-tech looking medical equipment, and clear signage. If the lobby of a hospital is dirty, you immediately assume the doctors are incompetent, even if they are world-class. The physical environment becomes the proxy for service quality.

The second characteristic is Perishability. Services cannot be stored for future sale or use. 

Let's say an APSRTC bus leaves Guntur for Visakhapatnam with ten empty seats. Can the bus depot store those ten empty seats and sell them tomorrow? No. Once the bus departs, that inventory, that potential revenue, is gone forever. This is why service industries are obsessed with capacity management and yield pricing—like hotels charging cheaper rates during the off-season to fill rooms.

The third characteristic is Inseparability. In product marketing, the factory makes the product, the retailer sells it, and the consumer eats it later. But in services, the service is produced and consumed at the very same time. 

Think about a classroom lecture. I, the teacher, am producing the service of education, and you, the student, are consuming it at the exact same moment. We co-create the value. If you are sleeping, the service fails. You cannot separate the provider from the service itself. This co-creation is crucial. In fitness coaching or weight-loss programs, if the client doesn't eat right, the service fails despite the trainer's best efforts.

.SCENE 3.2 - 10:30-15:30 [Slide 11]

INT. MEE SEVA CENTER / ARAKU TOUR CONGO - DAY

AUDIO:
Background: Continuing light instrumental.
Voice: Female narrator (continuing).
Sound Effects: Subtle transition sounds.

FEMALE NARRATOR
This leads us to the fourth characteristic: Variability. Services are highly variable because they depend on who provides them, when they are provided, and where they are provided. 

Let's use a funny, relatable example. You go to your favorite biryani point in Guntur on a Tuesday night. The biryani is absolutely delicious, perfectly spiced. You go back on Friday night, and it's too salty. Why? Because the chef was having a bad day, or there was a rush, or a different cook was on duty. Humans are not machines. We have moods, energy levels, and varying skills. Standardizing services to reduce this variability is one of the toughest challenges in business. That is why brands use detailed training manuals and standard operating procedures.

The fifth characteristic is Changing Demand. The demand for services is rarely stable. It fluctuates wildly based on seasons, days of the week, or even hours of the day. 

Take Araku Valley tourism. During winter and the holiday season, hotels are packed. During the hot summer, they are empty. Service marketers must design strategies to cope with these peaks and valleys—like offering special monsoon discounts or hiring temporary staff.

The sixth and final characteristic is Pricing. Since services are intangible and cannot be standardized, pricing is highly subjective. It is rarely based on a simple cost-plus formula. Instead, service pricing is determined by competition, perceived value, and demand. 

Think about a lawyer. Why does one lawyer charge one thousand rupees for an hour of advice, while another charges fifty thousand rupees? You aren't paying for their time; you are paying for their reputation, expertise, and the probability of winning your case.

Understanding these six characteristics is what separates rookie marketers from seasoned professionals.

.SCENE 4.1 - 15:30-19:00 [Slides 12-13]

INT. AIRTEL OFFICE / COACHING CENTER - DAY

AUDIO:
Background: Educational, clear instrumental (slightly academic).
Voice: Female narrator (systematic, easy to follow).
Sound Effects: Soft reveal sounds for each P.

FEMALE NARRATOR
Now that we understand the characteristics, let's look at how we market them. 

If you've studied basic marketing, you know the traditional four Ps: Product, Price, Place, and Promotion. But for services, the four Ps are simply not enough. We need to expand them to the Seven Ps. 

Let's break down the first four Ps in a service context. 

P1: Product. In services, the product is a bundle of intangible benefits. When you enroll in an IIT coaching center in Guntur, the product is not the physical classroom or the printed textbooks. The product is the transfer of knowledge, the problem-solving skills, and the career counseling. It is a transformational experience. In services, we also separate the core service, which is the education, from secondary services, like transport and online portal access.

P2: Price. Services pricing is highly dynamic. We see subscription models, where you pay monthly for Netflix or Airtel postpaid data packages. We see value-based pricing, where premium coaching centers charge higher fees because they have a track record of producing top rankers. The price must reflect the perceived value, because customers cannot inspect the service beforehand.

P3: Place. Place in services is all about accessibility and speed. In the digital age, place has moved from physical branches to our smartphones. Think about Swiggy and Zomato. They transformed the place utility of restaurants in Vijayawada and Guntur. The restaurant's kitchen is physical, but the distribution point is an app in your palm.

P4: Promotion. Because services are intangible, you cannot show pictures of the product. Instead, promotion must focus on tangibilizing the experience. This is why educational institutions do not just advertise their curriculum. They advertise pictures of smiling toppers, modern laboratory buildings, and lists of corporate recruiters. They sell the dream and show social proof. Word-of-mouth is crucial here; we trust our friends' dentist recommendations more than any billboard.

.SCENE 4.2 - 19:00-22:30 [Slides 12-13]

INT. NOVOTEL HOTEL / MEE SEVA CENTER - DAY

AUDIO:
Background: Continuing educational instrumental.
Voice: Female narrator (continuing).
Sound Effects: Subtle progress sounds.

FEMALE NARRATOR
Now, let's look at the three extra Ps that are absolutely crucial for services. 

P5: People. In a service business, people are the service. The customer's perception of your brand is completely determined by the behavior of your frontline employees. 

Imagine checking into the Novotel in Visakhapatnam. The rooms might be beautiful, the view of the sea spectacular. But if the receptionist at the front desk is rude, ignores you, and makes you wait, your entire experience is ruined. 

Here's an impromptu classroom insight: service employees perform "emotional labor." They must smile and be polite even when they are exhausted or dealing with a difficult client. This is why service brands invest heavily in employee training, motivation, and empowerment. Ritz-Carlton famously allows employees to spend up to two thousand dollars per day to resolve a guest complaint without manager approval. That is how much they value their people.

P6: Process. Process is the flow of activities and systems used to deliver the service. A bad process creates friction, and friction kills customer satisfaction. 

Think about going to a Mee Seva center to get a certificate. If the process requires you to fill out five paper forms, stand in three different lines, and wait for weeks, the service delivery is broken. Modern service companies design seamless, frictionless customer journeys. Amazon's one-click checkout is a classic example of process as a competitive advantage.

P7: Physical Evidence. As we discussed earlier under intangibility, physical evidence is the environment in which the service is delivered, along with any tangible elements that facilitate the service. 

Why do modern banks and corporate coaching offices in Guntur invest in fancy glass facades, air-conditioned lobbies, clean restrooms, and well-designed websites? Because these are physical proof of quality. 

Here's an impromptu design secret: did you know that hotels design olfactory environments? They create signature scents for their lobbies, and cafes use specific lighting and acoustic playlists to make you stay longer and buy more coffee. It is all physical evidence designed to impact your behavior. 

All seven Ps must align. You cannot have a premium price if your people are rude, your process is slow, and your physical evidence looks cheap. 

When you orchestrate these seven elements harmoniously, you create service excellence.

.SCENE 5.1 - 22:30-24:30 [Slide 1]

INT. MODERN UNIVERSITY LIBRARY - DAY

AUDIO:
Background: Same Indian classical instrumental as intro.
Voice: Professor's cloned voice (warm, encouraging).
Sound Effects: Subtle text reveal sounds.

PROFESSOR
Wonderful job! We have covered a massive amount of ground today in Video 2.

Let's do a quick recap:
We defined services marketing, looked at how Uber uses technology to build trust and credibility, and examined the six unique characteristics of services—intangibility, perishability, inseparability, variability, changing demand, and pricing. 

Finally, we explored the extended seven Ps framework, showing how Product, Price, Place, Promotion, People, Process, and Physical Evidence work together to define service quality in our local markets.

I hope you can now see why services marketing requires a completely different mindset compared to selling physical goods.

In our third and final video, we are going to step across borders. We will explore Global Marketing, examine the balance between standardization and adaptation, look at Nike and Domino's, and break down the eight massive challenges companies face when they decide to go global.

Until then, keep observing the services around you—whether you're ordering food, visiting a clinic, or boarding a bus. Think about how they apply the seven Ps.

Thank you, and Namaste!

TECHNICAL NOTES:
- Maintain consistency with intro segment
- Ensure smooth audio transition
- Professional closing animation
"""

# Define Video 3 in Fountain format: Global Marketing (Slides 14-21) - Expanded Classroom Style
VIDEO_3_FOUNTAIN = """Subject: Marketing Management
Title: Global Marketing
Page Length: Estimated ~17 pages
Duration: Around ~20 Minutes

.SCENE 1.1 - 00:00-01:30 [Slide 1]

INT. MODERN UNIVERSITY LIBRARY - DAY

AUDIO:
Background: Indian classical instrumental (same as Video 1).
Voice: Professor's cloned voice (warm, welcoming back).
Sound Effects: None (clean dialogue).

PROFESSOR
Welcome back, everyone! It is wonderful to have you here for the third and final installment of our Marketing Management series.

In Video 1, we established our local marketing foundations, and in Video 2, we mastered the intricacies of services marketing. We looked at examples close to home, analyzing Guntur chilli traders, Mangalagiri weavers, KIMS healthcare, and local transport. 

But today, we are going to expand our horizons. We are going to pack our bags and look at the global stage. 

What happens when a company decides that its domestic market is no longer big enough? What happens when a business wants to go global? 

In this video, we will explore Global Marketing. We will analyze the core strategic dilemma of standardization versus adaptation. We'll look at how Nike and Domino's Pizza have taken completely different paths to global dominance. And finally, we will examine the eight critical roadblocks and challenges that global managers must overcome.

Going global is exciting, but it is also filled with hidden landmines. So let's get ready, dive in, and explore the global marketplace.

TECHNICAL NOTES:
- Maintain visual consistency with Video 1
- Reference previous video naturally
- Set clear expectations for this video

> Smooth fade to first content segment.

.SCENE 2.1 - 01:30-06:00 [Slide 14]

EXT. GLOBAL CONNECTIVITY MAP - DAY

AUDIO:
Background: Cosmopolitan, worldly instrumental (subtle ethnic fusion).
Voice: Male narrator (sophisticated, global perspective).
Sound Effects: World map ambiance, subtle global soundscape.

MALE NARRATOR
Let's start with Slide 14. What exactly is Global Marketing? 

Put simply, global marketing is the process of planning, producing, placing, and promoting a company's products and services in a worldwide market. It goes far beyond mere exporting. It requires a company to understand the global environment and treat the world as its potential marketplace.

But as soon as a company decides to cross its borders, it is confronted with the single most important question in international business: 

Do we Standardize, or do we Adapt?

Standardization, or the global integration approach, is when a company sells the exact same product using the exact same marketing campaign worldwide. Think of Apple. An iPhone sold in New York is virtually identical to an iPhone sold in Guntur or Tokyo. The advertising campaigns look the same. 

Why do companies love standardization? Because it is incredibly cheap in the long run. It gives you massive economies of scale in manufacturing, cuts down on R&D costs, and creates a unified global brand image. 

On the other side of the coin, we have Adaptation, or local responsiveness. This is when a company customizes its product, packaging, pricing, and promotions to match the unique tastes, legal requirements, and cultural habits of each individual country. 

Here's an impromptu classroom anecdote on what happens when you suffer from home-country bias or "cultural myopia." Years ago, an American baby food company tried to sell its products in Africa. They used the same packaging design: a jar with a picture of a smiling baby on the label. 

But they didn't realize that in many of these local markets, the literacy rate was low, so the local custom was to put a picture of what was inside the jar on the label! The local consumers looked at the jar, saw the baby, and were absolutely horrified. They thought the jar contained baby meat! This is a classic example of why local adaptation and research are not optional. You cannot assume what works at home works everywhere.

This is why modern global companies prefer "Glocalization." Think globally, but act locally. Consider Starbucks. They standardize their coffee roasting process globally to maintain quality. But they adapt their menus and store designs locally. In India, they serve Cardamom Chai and Paneer wraps. In Japan, they sell Matcha lattes. They respect local tastes while keeping their global brand identity intact.

.SCENE 3.1 - 06:00-11:00 [Slides 15-16]

INT. RETAIL STORES - DAY

AUDIO:
Background: Continuing cosmopolitan instrumental.
Voice: Male narrator (continuing).
Sound Effects: Light transition sounds.

MALE NARRATOR
Let's look at Slide 15 and 16 to see how two world-class brands manage this balance.

First, let's look at Nike. 

Nike has built one of the most recognized brands on the planet. Their core product is athletic footwear and apparel. They rely heavily on standardization for their core products and their famous "Just Do It" slogan. 

However, Nike is a master of strategic international sponsorships and partnerships. In the United States, they sponsor American football and basketball. But when they entered Europe and South America, they put their resources behind soccer. When they came to India, what did they do? They sponsored the Indian national cricket team. 

Furthermore, Nike adapts its products technically. They design lightweight, breathable shoes for hot and humid tropical climates, and they engineered custom cricket shoes with specific spike configurations to grip Indian pitches. Even their ads are localized—their commercials in London feature grime music and street football, while in Mumbai they feature cricket played on crowded streets.

Nike understands that while their shoes are standard, the sport that captures the heart of the consumer is local. 

Now, let's contrast Nike with Domino's Pizza. 

Pizza is a global food, but taste is deeply personal. If you go to a Domino's in America, the menu is simple: pepperoni, sausage, cheese. 

But Domino's realized early on that if they wanted to win in international markets, they had to adapt their toppings, crusts, and sauces. 

If you visit a Domino's in Asia, you will find menu options like seafood and squid toppings in Japan, or bulgogi beef in South Korea. In India, they completely restructured their menu to offer a massive variety of vegetarian options, like Paneer Tikka pizza, Peppy Paneer, and spicy Keema toppings. 

But here's the process adaptation: Domino's couldn't just use standard delivery methods in traffic-congested Indian cities. They had to design custom hot bags and use small, maneuverable motorcycles to navigate the chaotic lanes of Mumbai, Delhi, and Bangalore to guarantee their thirty-minute delivery. 

Domino's proved that you cannot sell American taste preferences or delivery systems to the world. You must adapt your product to fit the local palate, while maintaining your fast delivery process.

.SCENE 4.1 - 11:00-17:00 [Slide 19]

INT. CORPORATE HQ / LEGAL OFFICE - DAY

AUDIO:
Background: Steady, slightly serious instrumental.
Voice: Male narrator (analytical, clear).
Sound Effects: Corporate atmosphere.

MALE NARRATOR
Now let's turn to Slide 19. Taking a brand global is not a walk in the park. There are eight critical challenges that companies must navigate. Let's look at the first four.

The first challenge is Cultural Differences. What appeals to a Western consumer might be completely offensive or irrelevant in another country. 

Consider McDonald's. When they entered India, they had a massive cultural hurdle: cows are sacred in India, and a large portion of the population is vegetarian. If McDonald's had insisted on selling their standard beef hamburgers, they would have gone bankrupt. Instead, they adapted. They removed beef from the menu, introduced the McAloo Tikki burger, and created a separate kitchen space for vegetarian food. 

Here's an impromptu insight: in the West, McDonald's is a cheap, fast convenience food for people on the go. In India, they positioned it as a premium family dining experience—clean, bright, and celebration-friendly. They adapted the entire brand perception. Similarly, beauty brands must adapt their cosmetic color palettes to match local skin tones and regional beauty standards—bronzers in the West versus dewy, glass-skin products in East Asia.

The second challenge is Legal and Regulatory Compliance. Every country has its own set of rules. 

In Europe, you have strict data privacy laws like GDPR. In Sweden, advertising toys to children under twelve on television is completely banned. In India, there are strict regulations regarding foreign direct investment and product labeling. If you don't comply, you face heavy fines or banishment.

The third challenge is Economic Uncertainty. Fluctuating exchange rates, inflation, and varying purchasing power make pricing and profit forecasting a nightmare. A product that is affordable in London might be a luxury item in Hanoi. 

Here's a clever adaptation: in emerging markets like India or the Philippines, consumer goods giants like Unilever sell shampoo and soap in small, single-use "sachets" for a few rupees. Consumers prefer small, daily cash outlays rather than buying bulk bottles. That is an economic adaptation.

The fourth challenge is Supply Chain Complexities. Managing logistics across continents involves dealing with different port infrastructures, customs regulations, tariffs, and transport networks. Establishing a reliable "cold chain" to transport temperature-sensitive items like ice cream or pharmaceuticals across regions with unstable electrical grids is a massive hurdle.

.SCENE 4.2 - 17:00-22:30 [Slide 20]

INT. GLOBAL OPERATIONS / HR CENTER - DAY

AUDIO:
Background: Continuing analytical instrumental.
Voice: Male narrator (continuing).
Sound Effects: Global communication sounds.

MALE NARRATOR
Let's turn to Slide 20 and look at the remaining four challenges of global marketing.

The fifth challenge is Political Risks. A sudden change in government, trade policies, or international relations can disrupt business overnight. 

If two countries enter a trade war, tariffs can skyrocket, making your imported goods uncompetitive. Governments can place sudden bans on foreign raw materials, forcing companies like automakers to source parts locally, or even ban foreign mobile applications.

The sixth challenge is Language Barriers. Communication is key to marketing, but direct translations can lead to embarrassing blunders. 

For instance, when KFC first entered China, their famous slogan "Finger-lickin' good" was translated into Chinese characters that literally meant "Eat your fingers off." 

Another famous example is the Chevy Nova. When General Motors launched the Nova in Central and South America, sales were terrible. Why? Because in Spanish, "No va" literally translates to "It doesn't go"! Who wants to buy a car called "It doesn't go"? Marketers must practice "transcreation"—translating the emotional feeling of a message rather than literal words.

The seventh challenge is Understanding Pricing Strategy. You cannot just convert your home currency prices. You must understand the local competitive landscape and consumer price sensitivity. 

This often leads to the problem of "Gray Markets" or parallel importing. For example, textbook publishers print cheap, softcover student editions for developing countries. But middle-men buy these cheap books in bulk and ship them copy to sell to students in the United States on eBay, undermining the publisher's high-margin domestic market.

The eighth and final challenge is Human Resource Management. Managing a diverse global team means dealing with different labor laws, work ethics, and cultural styles of leadership. 

A management style that works in Germany might be seen as overly aggressive in Thailand. In the US, executive decisions are often top-down and fast, whereas Japanese corporations use the "Ringi" system—a slow, bottom-up consensus-building process. Sometimes companies face high "expatriate failure rates" because they send managers abroad who don't understand these nuances. You must build a corporate culture that respects local differences.

Navigating these eight challenges is tough, but for those who succeed, the reward is a truly global brand.

> Fade to professor outro (1 second).

.SCENE 5.1 - 22:30-24:30 [Slide 1]

INT. MODERN UNIVERSITY LIBRARY - DAY

AUDIO:
Background: Same Indian classical instrumental as intro.
Voice: Professor's cloned voice (warm, encouraging).
Sound Effects: Subtle text reveal sounds.

PROFESSOR
Welcome back to the library. We have completed our journey through Global Marketing, and with that, our complete Unit-1 module on Marketing Management!

Let's do a final recap of Video 3:
We explored global marketing, analyzed the standardization versus adaptation dilemma, and looked at how Nike and Domino's customize their offerings for regional markets. 

Finally, we detailed the eight critical challenges of global marketing—cultural differences, legal compliance, economic factors, supply chain issues, political risks, language barriers, pricing, and human resources.

By combining the local market foundations from Video 1, the services dynamics from Video 2, and the global perspectives from Video 3, you now have a powerful, comprehensive toolkit to analyze any marketing problem.

Remember, marketing is not just about memorizing theories. It is about keeping your eyes open, observing consumer behavior, and understanding the world around you.

Thank you so much for your time, your focus, and your curiosity. Keep learning, keep questioning, and keep exploring.

Namaste!

TECHNICAL NOTES:
- Maintain consistency with intro segment
- Ensure smooth audio transition
- Professional closing animation
"""

def generate_fountain_files():
    base_dir = "/Users/rkanduk/Documents/GitHub/OpenMontage"
    
    # Save the Fountain format scripts
    f1_path = os.path.join(base_dir, "Video_1_Foundations_of_Marketing.fountain")
    with open(f1_path, "w", encoding="utf-8") as f:
        f.write(VIDEO_1_FOUNTAIN)
    print(f"Saved Fountain script: {f1_path}")
        
    f2_path = os.path.join(base_dir, "Video_2_Services_Marketing.fountain")
    with open(f2_path, "w", encoding="utf-8") as f:
        f.write(VIDEO_2_FOUNTAIN)
    print(f"Saved Fountain script: {f2_path}")
    
    f3_path = os.path.join(base_dir, "Video_3_Global_Marketing.fountain")
    with open(f3_path, "w", encoding="utf-8") as f:
        f.write(VIDEO_3_FOUNTAIN)
    print(f"Saved Fountain script: {f3_path}")
    
    return f1_path, f2_path, f3_path

def set_font(run, name="Courier New", size=12, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def convert_fountain_to_docx(fountain_path, docx_path):
    # Parse fountain script using the installed fountain library
    with open(fountain_path, "r", encoding="utf-8") as f:
        script_text = f.read()
        
    fountain_parser = Fountain()
    tokens = fountain_parser.tokenize(script_text.splitlines())
    
    doc = Document()
    
    # Configure margins for standard screenplay look
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1)
        
    # We will write the titles manually first to ensure exact metadata ordering
    title_keys = ["subject", "title", "page length", "duration"]
    title_dict = {}
    
    # Extract titles token if exists
    for token in tokens:
        if token.get("type") == "titles":
            title_dict = token.get("title-fields", {})
            break
            
    # Output the metadata ordered correctly
    if title_dict:
        for k in title_keys:
            val = title_dict.get(k)
            if val:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                # Capitalize key
                k_cap = " ".join([w.capitalize() for w in k.split()])
                run = p.add_run(f"{k_cap}: {val}")
                set_font(run, bold=True)
        doc.add_paragraph() # extra spacer
        
    for token in tokens:
        ttype = token.get("type")
        lines = token.get("lines", [])
        
        if not lines:
            continue
            
        text = "".join(lines).strip()
        
        # Skip outputting titles again
        if ttype == "titles":
            continue
            
        elif ttype == "slugline":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(text)
            set_font(run, bold=True)
            
        elif ttype == "dialogue":
            # Character name
            char_p = doc.add_paragraph()
            char_p.paragraph_format.left_indent = Inches(2.2)
            char_p.paragraph_format.space_before = Pt(12)
            char_p.paragraph_format.space_after = Pt(0)
            run_char = char_p.add_run(token.get("character", ""))
            set_font(run_char, bold=True)
            
            # Dialogue lines
            dialogue_lines = token.get("text", [])
            for line in dialogue_lines:
                # Check for parenthetical
                if line.startswith("(") and line.endswith(")"):
                    p_para = doc.add_paragraph()
                    p_para.paragraph_format.left_indent = Inches(1.6)
                    p_para.paragraph_format.space_after = Pt(0)
                    run_p = p_para.add_run(line)
                    set_font(run_p, italic=True)
                else:
                    d_para = doc.add_paragraph()
                    d_para.paragraph_format.left_indent = Inches(1.0)
                    d_para.paragraph_format.right_margin = Inches(1.5)
                    d_para.paragraph_format.space_after = Pt(12)
                    run_d = d_para.add_run(line)
                    set_font(run_d)
                    
        elif ttype == "transition":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(3.5)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(text)
            set_font(run, italic=True)
            
        elif ttype == "action":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            
            # Format visual/audio cues nicely
            if text.startswith("AUDIO:") or text.startswith("VISUAL:") or text.startswith("Visual shows"):
                run = p.add_run(text)
                set_font(run, size=11, italic=True)
            else:
                run = p.add_run(text)
                set_font(run)
                
    doc.save(docx_path)
    print(f"Exported {fountain_path} to Word Doc: {docx_path}")

if __name__ == "__main__":
    f1, f2, f3 = generate_fountain_files()
    convert_fountain_to_docx(f1, f1.replace(".fountain", ".docx"))
    convert_fountain_to_docx(f2, f2.replace(".fountain", ".docx"))
    convert_fountain_to_docx(f3, f3.replace(".fountain", ".docx"))
